"""
Report Generator Service - Generates comprehensive business analysis reports
Provides PDF, DOCX, and JSON export with professional formatting and charts
"""
import logging
import io
import json
import hashlib
import secrets
from typing import Dict, List, Any, Optional, Union
from datetime import datetime, timezone, timedelta
from io import BytesIO
from pydantic import BaseModel, Field

# ReportLab imports for PDF generation
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.colors import HexColor
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, 
    PageBreak, Image as RLImage, KeepTogether
)
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader

# Matplotlib for chart generation
import matplotlib
matplotlib.use('Agg')  # Use non-interactive backend
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.figure import Figure

# Python-docx for DOCX generation
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Import data models from other services
from budget_analyzer import BudgetBreakdown, CategoryAllocation, SubCategory
from financial_projector import (
    FinancialProjections, RevenueProjection, CostProjection,
    BreakEvenAnalysis, ProfitMargins, ROICalculation, CashFlowProjection
)
from risk_assessment import RiskAssessmentResult, Risk
from growth_strategy import GrowthStrategyResult, GrowthPhase, ActionItem

logger = logging.getLogger(__name__)


class Report(BaseModel):
    """Report metadata"""
    report_id: str
    session_id: str
    user_id: str
    format: str  # "pdf", "docx", "json"
    file_path: Optional[str] = None
    file_size: Optional[int] = None
    generation_timestamp: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))


class SharingLink(BaseModel):
    """Secure sharing link"""
    link_id: str
    report_id: str
    url: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    expires_at: datetime
    access_count: int = 0
    last_accessed_at: Optional[datetime] = None


class AnalysisData(BaseModel):
    """Complete analysis data for report generation"""
    session_id: str
    user_id: str
    business_idea: Dict[str, Any]
    budget: BudgetBreakdown
    market_research: Optional[Dict[str, Any]] = None
    financial_projections: FinancialProjections
    risk_assessment: RiskAssessmentResult
    growth_strategy: GrowthStrategyResult
    generation_timestamp: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))


class ReportGenerator:
    """
    Report generator that compiles analysis data into professional reports
    Supports PDF, DOCX, and JSON formats with charts and visualizations
    """
    
    # AstraMark brand colors
    BRAND_COLORS = {
        "primary": HexColor("#8B5CF6"),      # Purple
        "secondary": HexColor("#EC4899"),    # Pink
        "accent": HexColor("#10B981"),       # Green
        "warning": HexColor("#F59E0B"),      # Orange
        "danger": HexColor("#EF4444"),       # Red
        "info": HexColor("#3B82F6"),         # Blue
        "dark": HexColor("#1F2937"),         # Dark gray
        "light": HexColor("#F3F4F6"),        # Light gray
        "white": HexColor("#FFFFFF")
    }
    
    def __init__(self, db=None):
        """
        Initialize report generator
        
        Args:
            db: MongoDB database connection (optional)
        """
        self.db = db
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()
        logger.info("ReportGenerator initialized")
    
    def _setup_custom_styles(self):
        """Setup custom paragraph styles for PDF"""
        # Cover page title
        self.styles.add(ParagraphStyle(
            name='CoverTitle',
            parent=self.styles['Heading1'],
            fontSize=32,
            textColor=self.BRAND_COLORS["primary"],
            spaceAfter=20,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        ))
        
        # Section header
        self.styles.add(ParagraphStyle(
            name='SectionHeader',
            parent=self.styles['Heading1'],
            fontSize=18,
            textColor=self.BRAND_COLORS["primary"],
            spaceBefore=20,
            spaceAfter=12,
            fontName='Helvetica-Bold'
        ))
        
        # Subsection header
        self.styles.add(ParagraphStyle(
            name='SubsectionHeader',
            parent=self.styles['Heading2'],
            fontSize=14,
            textColor=self.BRAND_COLORS["secondary"],
            spaceBefore=15,
            spaceAfter=10,
            fontName='Helvetica-Bold'
        ))
        
        # Body text
        self.styles.add(ParagraphStyle(
            name='CustomBody',
            parent=self.styles['BodyText'],
            fontSize=11,
            spaceAfter=12,
            alignment=TA_JUSTIFY
        ))
        
        # Bullet point
        self.styles.add(ParagraphStyle(
            name='BulletPoint',
            parent=self.styles['BodyText'],
            fontSize=11,
            leftIndent=20,
            spaceAfter=6
        ))

    async def generate_report(
        self,
        analysis_data: AnalysisData,
        format: str = "pdf"
    ) -> Report:
        """
        Generate complete feasibility report
        
        Args:
            analysis_data: Complete analysis data
            format: Report format ("pdf", "docx", "json")
            
        Returns:
            Report metadata
        """
        logger.info(f"Generating {format.upper()} report for session {analysis_data.session_id}")
        
        report_id = self._generate_report_id()
        
        if format == "pdf":
            file_buffer = await self.generate_pdf(analysis_data)
        elif format == "docx":
            file_buffer = await self.generate_docx(analysis_data)
        elif format == "json":
            json_data = await self.generate_json(analysis_data)
            file_buffer = BytesIO(json.dumps(json_data, indent=2, default=str).encode('utf-8'))
        else:
            raise ValueError(f"Unsupported format: {format}")
        
        # Get file size
        file_buffer.seek(0, 2)  # Seek to end
        file_size = file_buffer.tell()
        file_buffer.seek(0)  # Reset to beginning
        
        report = Report(
            report_id=report_id,
            session_id=analysis_data.session_id,
            user_id=analysis_data.user_id,
            format=format,
            file_size=file_size
        )
        
        logger.info(f"Report generated: {report_id} ({file_size} bytes)")
        return report, file_buffer
    
    async def generate_pdf(self, analysis_data: AnalysisData) -> BytesIO:
        """
        Generate PDF report with charts and tables
        
        Args:
            analysis_data: Complete analysis data
            
        Returns:
            BytesIO buffer containing PDF
        """
        logger.info("Generating PDF report")
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        story = []
        
        # Cover Page
        story.extend(self._create_cover_page(analysis_data))
        story.append(PageBreak())
        
        # Table of Contents
        story.extend(self._create_table_of_contents())
        story.append(PageBreak())
        
        # Executive Summary
        story.extend(self._create_executive_summary(analysis_data))
        story.append(PageBreak())
        
        # Business Overview
        story.extend(self._create_business_overview(analysis_data))
        story.append(PageBreak())
        
        # Market Analysis
        if analysis_data.market_research:
            story.extend(self._create_market_analysis(analysis_data))
            story.append(PageBreak())
        
        # Financial Projections
        story.extend(self._create_financial_projections_section(analysis_data))
        story.append(PageBreak())
        
        # Budget Breakdown
        story.extend(self._create_budget_breakdown_section(analysis_data))
        story.append(PageBreak())
        
        # Risk Assessment
        story.extend(self._create_risk_assessment_section(analysis_data))
        story.append(PageBreak())
        
        # Growth Strategy
        story.extend(self._create_growth_strategy_section(analysis_data))
        story.append(PageBreak())
        
        # Appendices
        story.extend(self._create_appendices(analysis_data))
        
        # Build PDF
        doc.build(story, onFirstPage=self._add_page_number, onLaterPages=self._add_page_number)
        
        buffer.seek(0)
        logger.info("PDF report generated successfully")
        return buffer
    
    async def generate_docx(self, analysis_data: AnalysisData) -> BytesIO:
        """
        Generate DOCX report for editing
        
        Args:
            analysis_data: Complete analysis data
            
        Returns:
            BytesIO buffer containing DOCX
        """
        logger.info("Generating DOCX report")
        
        doc = Document()
        
        # Set document styles
        self._setup_docx_styles(doc)
        
        # Cover Page
        self._add_docx_cover_page(doc, analysis_data)
        doc.add_page_break()
        
        # Executive Summary
        self._add_docx_executive_summary(doc, analysis_data)
        doc.add_page_break()
        
        # Business Overview
        self._add_docx_business_overview(doc, analysis_data)
        doc.add_page_break()
        
        # Market Analysis
        if analysis_data.market_research:
            self._add_docx_market_analysis(doc, analysis_data)
            doc.add_page_break()
        
        # Financial Projections
        self._add_docx_financial_projections(doc, analysis_data)
        doc.add_page_break()
        
        # Budget Breakdown
        self._add_docx_budget_breakdown(doc, analysis_data)
        doc.add_page_break()
        
        # Risk Assessment
        self._add_docx_risk_assessment(doc, analysis_data)
        doc.add_page_break()
        
        # Growth Strategy
        self._add_docx_growth_strategy(doc, analysis_data)
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        logger.info("DOCX report generated successfully")
        return buffer
    
    async def generate_json(self, analysis_data: AnalysisData) -> Dict[str, Any]:
        """
        Generate JSON export for structured data
        
        Args:
            analysis_data: Complete analysis data
            
        Returns:
            Dictionary with structured data
        """
        logger.info("Generating JSON export")
        
        json_data = {
            "report_metadata": {
                "session_id": analysis_data.session_id,
                "user_id": analysis_data.user_id,
                "generation_timestamp": analysis_data.generation_timestamp.isoformat(),
                "format": "json",
                "version": "1.0"
            },
            "business_idea": analysis_data.business_idea,
            "budget": analysis_data.budget.dict(),
            "market_research": analysis_data.market_research,
            "financial_projections": analysis_data.financial_projections.dict(),
            "risk_assessment": analysis_data.risk_assessment.dict(),
            "growth_strategy": analysis_data.growth_strategy.dict()
        }
        
        logger.info("JSON export generated successfully")
        return json_data

    # ==================== CHART GENERATION METHODS ====================
    
    def create_revenue_chart(self, projections: List[RevenueProjection]) -> bytes:
        """
        Generate revenue projection chart
        
        Args:
            projections: List of revenue projections
            
        Returns:
            PNG image bytes
        """
        logger.debug("Creating revenue projection chart")
        
        fig, ax = plt.subplots(figsize=(10, 6))
        
        # Extract data
        periods = [p.period for p in projections]
        revenues = [p.revenue for p in projections]
        
        # Create line chart
        ax.plot(periods, revenues, marker='o', linewidth=2, 
                color='#8B5CF6', markersize=6)
        
        # Styling
        ax.set_xlabel('Period', fontsize=12, fontweight='bold')
        ax.set_ylabel('Revenue', fontsize=12, fontweight='bold')
        ax.set_title('Revenue Projections', fontsize=14, fontweight='bold', pad=20)
        ax.grid(True, alpha=0.3)
        
        # Rotate x-axis labels for readability
        plt.xticks(rotation=45, ha='right')
        
        # Format y-axis as currency
        ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'${x:,.0f}'))
        
        plt.tight_layout()
        
        # Save to bytes
        buffer = BytesIO()
        plt.savefig(buffer, format='png', dpi=150, bbox_inches='tight')
        plt.close(fig)
        buffer.seek(0)
        
        return buffer.getvalue()
    
    def create_budget_pie_chart(self, breakdown: BudgetBreakdown) -> bytes:
        """
        Generate budget allocation pie chart
        
        Args:
            breakdown: Budget breakdown data
            
        Returns:
            PNG image bytes
        """
        logger.debug("Creating budget allocation pie chart")
        
        fig, ax = plt.subplots(figsize=(10, 8))
        
        # Extract data
        categories = [c.category for c in breakdown.categories]
        amounts = [c.amount for c in breakdown.categories]
        
        # Colors matching brand
        colors = ['#8B5CF6', '#EC4899', '#10B981', '#F59E0B', '#3B82F6']
        
        # Create pie chart
        wedges, texts, autotexts = ax.pie(
            amounts,
            labels=categories,
            autopct='%1.1f%%',
            startangle=90,
            colors=colors,
            textprops={'fontsize': 11}
        )
        
        # Make percentage text bold and white
        for autotext in autotexts:
            autotext.set_color('white')
            autotext.set_fontweight('bold')
        
        ax.set_title('Budget Allocation by Category', 
                     fontsize=14, fontweight='bold', pad=20)
        
        plt.tight_layout()
        
        # Save to bytes
        buffer = BytesIO()
        plt.savefig(buffer, format='png', dpi=150, bbox_inches='tight')
        plt.close(fig)
        buffer.seek(0)
        
        return buffer.getvalue()
    
    def create_cash_flow_chart(self, cash_flow: List[CashFlowProjection]) -> bytes:
        """
        Generate cash flow chart
        
        Args:
            cash_flow: List of cash flow projections
            
        Returns:
            PNG image bytes
        """
        logger.debug("Creating cash flow chart")
        
        fig, ax = plt.subplots(figsize=(12, 6))
        
        # Extract data (limit to first 12 months for readability)
        periods = [cf.period for cf in cash_flow[:12]]
        inflows = [cf.inflows for cf in cash_flow[:12]]
        outflows = [cf.outflows for cf in cash_flow[:12]]
        net_cash = [cf.net_cash_flow for cf in cash_flow[:12]]
        
        # Create bar chart
        x = range(len(periods))
        width = 0.25
        
        ax.bar([i - width for i in x], inflows, width, 
               label='Inflows', color='#10B981', alpha=0.8)
        ax.bar(x, outflows, width, 
               label='Outflows', color='#EF4444', alpha=0.8)
        ax.bar([i + width for i in x], net_cash, width, 
               label='Net Cash Flow', color='#8B5CF6', alpha=0.8)
        
        # Styling
        ax.set_xlabel('Period', fontsize=12, fontweight='bold')
        ax.set_ylabel('Amount', fontsize=12, fontweight='bold')
        ax.set_title('Cash Flow Projections (First 12 Months)', 
                     fontsize=14, fontweight='bold', pad=20)
        ax.set_xticks(x)
        ax.set_xticklabels(periods, rotation=45, ha='right')
        ax.legend(loc='upper left')
        ax.grid(True, alpha=0.3, axis='y')
        
        # Format y-axis as currency
        ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'${x:,.0f}'))
        
        # Add zero line
        ax.axhline(y=0, color='black', linestyle='-', linewidth=0.5)
        
        plt.tight_layout()
        
        # Save to bytes
        buffer = BytesIO()
        plt.savefig(buffer, format='png', dpi=150, bbox_inches='tight')
        plt.close(fig)
        buffer.seek(0)
        
        return buffer.getvalue()

    # ==================== SHARING LINK METHODS ====================
    
    async def create_sharing_link(
        self,
        report_id: str,
        expiration_days: int = 30
    ) -> SharingLink:
        """
        Generate secure sharing link with expiration
        
        Args:
            report_id: Report ID to share
            expiration_days: Days until link expires (7-90)
            
        Returns:
            SharingLink object
        """
        logger.info(f"Creating sharing link for report {report_id}")
        
        # Validate expiration days
        if expiration_days < 7 or expiration_days > 90:
            raise ValueError("Expiration days must be between 7 and 90")
        
        # Generate secure link ID
        link_id = secrets.token_urlsafe(32)
        
        # Calculate expiration
        expires_at = datetime.now(timezone.utc) + timedelta(days=expiration_days)
        
        # Generate URL (placeholder - would be actual domain in production)
        url = f"https://astramark.com/shared/reports/{link_id}"
        
        sharing_link = SharingLink(
            link_id=link_id,
            report_id=report_id,
            url=url,
            expires_at=expires_at
        )
        
        # Store in database if available
        if self.db:
            await self.db.sharing_links.insert_one(sharing_link.dict())
        
        logger.info(f"Sharing link created: {link_id} (expires {expires_at})")
        return sharing_link
    
    async def track_share_access(
        self,
        link_id: str,
        accessor_info: Dict[str, Any]
    ) -> None:
        """
        Track sharing link access
        
        Args:
            link_id: Sharing link ID
            accessor_info: Information about accessor (IP, user agent, etc.)
        """
        logger.info(f"Tracking access to sharing link {link_id}")
        
        if self.db:
            await self.db.sharing_links.update_one(
                {"link_id": link_id},
                {
                    "$inc": {"access_count": 1},
                    "$set": {
                        "last_accessed_at": datetime.now(timezone.utc)
                    },
                    "$push": {
                        "access_log": {
                            "timestamp": datetime.now(timezone.utc),
                            "accessor_info": accessor_info
                        }
                    }
                }
            )
    
    # ==================== PDF SECTION CREATION METHODS ====================
    
    def _create_cover_page(self, analysis_data: AnalysisData) -> List:
        """Create PDF cover page"""
        elements = []
        
        # Title
        elements.append(Spacer(1, 2*inch))
        elements.append(Paragraph(
            "Business Feasibility Report",
            self.styles['CoverTitle']
        ))
        
        elements.append(Spacer(1, 0.5*inch))
        
        # Business name/title
        business_title = analysis_data.business_idea.get('description', 'Business Analysis')[:100]
        elements.append(Paragraph(
            business_title,
            self.styles['Heading2']
        ))
        
        elements.append(Spacer(1, 1*inch))
        
        # Metadata
        metadata_style = ParagraphStyle(
            'Metadata',
            parent=self.styles['Normal'],
            fontSize=12,
            alignment=TA_CENTER
        )
        
        elements.append(Paragraph(
            f"<b>Report ID:</b> {analysis_data.session_id[:16]}",
            metadata_style
        ))
        elements.append(Spacer(1, 0.2*inch))
        
        elements.append(Paragraph(
            f"<b>Generated:</b> {analysis_data.generation_timestamp.strftime('%B %d, %Y')}",
            metadata_style
        ))
        elements.append(Spacer(1, 0.2*inch))
        
        # Budget
        currency_symbol = self._get_currency_symbol(analysis_data.budget.currency)
        formatted_budget = f"{currency_symbol}{analysis_data.budget.total_budget:,.0f}"
        elements.append(Paragraph(
            f"<b>Budget:</b> {formatted_budget}",
            metadata_style
        ))
        
        elements.append(Spacer(1, 2*inch))
        
        # Footer
        elements.append(Paragraph(
            "Generated by AstraMark AI Business Analysis Platform",
            metadata_style
        ))
        
        return elements
    
    def _create_table_of_contents(self) -> List:
        """Create table of contents"""
        elements = []
        
        elements.append(Paragraph("Table of Contents", self.styles['SectionHeader']))
        elements.append(Spacer(1, 0.3*inch))
        
        toc_items = [
            "1. Executive Summary",
            "2. Business Overview",
            "3. Market Analysis",
            "4. Financial Projections",
            "5. Budget Breakdown",
            "6. Risk Assessment",
            "7. Growth Strategy",
            "8. Appendices"
        ]
        
        for item in toc_items:
            elements.append(Paragraph(item, self.styles['CustomBody']))
            elements.append(Spacer(1, 0.1*inch))
        
        return elements
    
    def _create_executive_summary(self, analysis_data: AnalysisData) -> List:
        """Create executive summary section"""
        elements = []
        
        elements.append(Paragraph("Executive Summary", self.styles['SectionHeader']))
        elements.append(Spacer(1, 0.2*inch))
        
        # Business overview
        business_desc = analysis_data.business_idea.get('description', 'N/A')
        elements.append(Paragraph(
            f"<b>Business Concept:</b> {business_desc}",
            self.styles['CustomBody']
        ))
        elements.append(Spacer(1, 0.1*inch))
        
        # Key findings
        elements.append(Paragraph("Key Findings", self.styles['SubsectionHeader']))
        
        # Financial highlights
        fp = analysis_data.financial_projections
        year1_revenue = sum(rp.revenue for rp in fp.revenue_projections[:12])
        
        findings = [
            f"Projected Year 1 Revenue: ${year1_revenue:,.0f}",
            f"Break-even Point: Month {fp.break_even_analysis.break_even_month}",
            f"3-Year ROI: {fp.roi_calculations['3_year'].roi_percentage:.1f}%",
            f"Total Risks Identified: {analysis_data.risk_assessment.total_risk_count} "
            f"({analysis_data.risk_assessment.high_risk_count} high priority)",
            f"Growth Phases: {len(analysis_data.growth_strategy.growth_phases)} phases over 36 months"
        ]
        
        for finding in findings:
            elements.append(Paragraph(f"• {finding}", self.styles['BulletPoint']))
        
        elements.append(Spacer(1, 0.2*inch))
        
        # Recommendation
        elements.append(Paragraph("Recommendation", self.styles['SubsectionHeader']))
        
        # Simple recommendation based on ROI and break-even
        roi_3yr = fp.roi_calculations['3_year'].roi_percentage
        break_even_month = fp.break_even_analysis.break_even_month
        
        if roi_3yr > 100 and break_even_month <= 18:
            recommendation = "HIGHLY RECOMMENDED: Strong financial projections with excellent ROI and reasonable break-even timeline."
        elif roi_3yr > 50 and break_even_month <= 24:
            recommendation = "RECOMMENDED: Solid business opportunity with good ROI potential and manageable risk profile."
        elif roi_3yr > 0:
            recommendation = "PROCEED WITH CAUTION: Positive ROI but requires careful execution and risk management."
        else:
            recommendation = "REQUIRES REVISION: Current projections suggest challenges. Consider refining business model or budget allocation."
        
        elements.append(Paragraph(recommendation, self.styles['CustomBody']))
        
        return elements

    def _create_business_overview(self, analysis_data: AnalysisData) -> List:
        """Create business overview section"""
        elements = []
        
        elements.append(Paragraph("Business Overview", self.styles['SectionHeader']))
        elements.append(Spacer(1, 0.2*inch))
        
        # Business idea details
        idea = analysis_data.business_idea
        
        details = [
            ("Description", idea.get('description', 'N/A')),
            ("Industry", idea.get('industry', 'N/A')),
            ("Target Market", idea.get('target_market', 'N/A')),
            ("Geographic Location", idea.get('geographic_location', 'N/A')),
            ("Product/Service Type", idea.get('product_service_type', 'N/A'))
        ]
        
        for label, value in details:
            elements.append(Paragraph(
                f"<b>{label}:</b> {value}",
                self.styles['CustomBody']
            ))
        
        elements.append(Spacer(1, 0.2*inch))
        
        # Budget summary
        elements.append(Paragraph("Budget Summary", self.styles['SubsectionHeader']))
        
        currency_symbol = self._get_currency_symbol(analysis_data.budget.currency)
        elements.append(Paragraph(
            f"<b>Total Investment Budget:</b> {currency_symbol}{analysis_data.budget.total_budget:,.0f}",
            self.styles['CustomBody']
        ))
        
        return elements
    
    def _create_market_analysis(self, analysis_data: AnalysisData) -> List:
        """Create market analysis section"""
        elements = []
        
        elements.append(Paragraph("Market Analysis", self.styles['SectionHeader']))
        elements.append(Spacer(1, 0.2*inch))
        
        market = analysis_data.market_research or {}
        
        # Market size
        if 'market_size' in market:
            elements.append(Paragraph("Market Size", self.styles['SubsectionHeader']))
            elements.append(Paragraph(
                f"Total Addressable Market: {market['market_size'].get('value', 'N/A')}",
                self.styles['CustomBody']
            ))
            elements.append(Spacer(1, 0.1*inch))
        
        # Competitors
        if 'competitors' in market:
            elements.append(Paragraph("Competitive Landscape", self.styles['SubsectionHeader']))
            elements.append(Paragraph(
                f"Identified {len(market['competitors'])} direct competitors",
                self.styles['CustomBody']
            ))
            elements.append(Spacer(1, 0.1*inch))
            
            # Competitor table (top 5)
            comp_data = [['Competitor', 'Strengths', 'Weaknesses']]
            for comp in market['competitors'][:5]:
                comp_data.append([
                    comp.get('name', 'N/A'),
                    ', '.join(comp.get('strengths', [])[:2]),
                    ', '.join(comp.get('weaknesses', [])[:2])
                ])
            
            comp_table = Table(comp_data, colWidths=[2*inch, 2.5*inch, 2.5*inch])
            comp_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), self.BRAND_COLORS["primary"]),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 11),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), self.BRAND_COLORS["light"]),
                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                ('VALIGN', (0, 0), (-1, -1), 'TOP')
            ]))
            elements.append(comp_table)
        
        # Trends
        if 'trends' in market:
            elements.append(Spacer(1, 0.2*inch))
            elements.append(Paragraph("Industry Trends", self.styles['SubsectionHeader']))
            
            for trend in market['trends'][:5]:
                elements.append(Paragraph(
                    f"• <b>{trend.get('title', 'N/A')}:</b> {trend.get('description', 'N/A')}",
                    self.styles['BulletPoint']
                ))
        
        return elements
    
    def _create_financial_projections_section(self, analysis_data: AnalysisData) -> List:
        """Create financial projections section"""
        elements = []
        
        elements.append(Paragraph("Financial Projections", self.styles['SectionHeader']))
        elements.append(Spacer(1, 0.2*inch))
        
        fp = analysis_data.financial_projections
        
        # Revenue chart
        elements.append(Paragraph("Revenue Forecast", self.styles['SubsectionHeader']))
        revenue_chart_bytes = self.create_revenue_chart(fp.revenue_projections)
        revenue_img = RLImage(BytesIO(revenue_chart_bytes), width=6*inch, height=3.6*inch)
        elements.append(revenue_img)
        elements.append(Spacer(1, 0.2*inch))
        
        # Break-even analysis
        elements.append(Paragraph("Break-Even Analysis", self.styles['SubsectionHeader']))
        be = fp.break_even_analysis
        
        be_data = [
            ['Metric', 'Value'],
            ['Break-Even Month', f"Month {be.break_even_month}"],
            ['Break-Even Revenue', f"${be.break_even_revenue:,.0f}"],
            ['Cumulative Investment', f"${be.cumulative_investment_at_break_even:,.0f}"],
            ['Months to Profitability', f"{be.months_to_profitability} months"]
        ]
        
        be_table = Table(be_data, colWidths=[3*inch, 3*inch])
        be_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), self.BRAND_COLORS["info"]),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
            ('BACKGROUND', (0, 1), (-1, -1), self.BRAND_COLORS["light"])
        ]))
        elements.append(be_table)
        elements.append(Spacer(1, 0.2*inch))
        
        # Profit margins
        elements.append(Paragraph("Profit Margins (Year 1 Average)", self.styles['SubsectionHeader']))
        pm = fp.profit_margins
        
        pm_data = [
            ['Margin Type', 'Percentage'],
            ['Gross Margin', f"{pm.gross_margin:.1f}%"],
            ['Operating Margin', f"{pm.operating_margin:.1f}%"],
            ['Net Margin', f"{pm.net_margin:.1f}%"]
        ]
        
        pm_table = Table(pm_data, colWidths=[3*inch, 3*inch])
        pm_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), self.BRAND_COLORS["accent"]),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
            ('BACKGROUND', (0, 1), (-1, -1), self.BRAND_COLORS["light"])
        ]))
        elements.append(pm_table)
        elements.append(Spacer(1, 0.2*inch))
        
        # ROI calculations
        elements.append(Paragraph("Return on Investment (ROI)", self.styles['SubsectionHeader']))
        
        roi_data = [['Period', 'ROI %', 'Total Return', 'Payback Period']]
        for period_key, roi in fp.roi_calculations.items():
            period_label = period_key.replace('_', ' ').title()
            roi_data.append([
                period_label,
                f"{roi.roi_percentage:.1f}%",
                f"${roi.total_return:,.0f}",
                f"{roi.payback_period_months} months"
            ])
        
        roi_table = Table(roi_data, colWidths=[1.5*inch, 1.5*inch, 2*inch, 2*inch])
        roi_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), self.BRAND_COLORS["secondary"]),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
            ('BACKGROUND', (0, 1), (-1, -1), self.BRAND_COLORS["light"])
        ]))
        elements.append(roi_table)
        elements.append(Spacer(1, 0.2*inch))
        
        # Cash flow chart
        elements.append(Paragraph("Cash Flow Projections", self.styles['SubsectionHeader']))
        cash_flow_chart_bytes = self.create_cash_flow_chart(fp.cash_flow_projections)
        cash_flow_img = RLImage(BytesIO(cash_flow_chart_bytes), width=6*inch, height=3*inch)
        elements.append(cash_flow_img)
        
        return elements

    def _create_budget_breakdown_section(self, analysis_data: AnalysisData) -> List:
        """Create budget breakdown section"""
        elements = []
        
        elements.append(Paragraph("Budget Breakdown", self.styles['SectionHeader']))
        elements.append(Spacer(1, 0.2*inch))
        
        budget = analysis_data.budget
        
        # Budget pie chart
        elements.append(Paragraph("Budget Allocation", self.styles['SubsectionHeader']))
        budget_chart_bytes = self.create_budget_pie_chart(budget)
        budget_img = RLImage(BytesIO(budget_chart_bytes), width=5*inch, height=4*inch)
        elements.append(budget_img)
        elements.append(Spacer(1, 0.2*inch))
        
        # Category breakdown table
        elements.append(Paragraph("Category Details", self.styles['SubsectionHeader']))
        
        currency_symbol = self._get_currency_symbol(budget.currency)
        
        cat_data = [['Category', 'Amount', 'Percentage', 'Rationale']]
        for cat in budget.categories:
            cat_data.append([
                cat.category,
                f"{currency_symbol}{cat.amount:,.0f}",
                f"{cat.percentage:.1f}%",
                cat.rationale[:50] + "..." if len(cat.rationale) > 50 else cat.rationale
            ])
        
        cat_table = Table(cat_data, colWidths=[1.5*inch, 1.5*inch, 1*inch, 3*inch])
        cat_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), self.BRAND_COLORS["primary"]),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
            ('BACKGROUND', (0, 1), (-1, -1), self.BRAND_COLORS["light"]),
            ('VALIGN', (0, 0), (-1, -1), 'TOP')
        ]))
        elements.append(cat_table)
        elements.append(Spacer(1, 0.2*inch))
        
        # Sub-categories for each major category
        elements.append(Paragraph("Detailed Sub-Categories", self.styles['SubsectionHeader']))
        
        for category_name, sub_cats in budget.sub_categories.items():
            elements.append(Paragraph(f"<b>{category_name}</b>", self.styles['CustomBody']))
            
            sub_data = [['Item', 'Amount', 'Priority']]
            for sub in sub_cats:
                sub_data.append([
                    sub.name,
                    f"{currency_symbol}{sub.amount:,.0f}",
                    sub.priority
                ])
            
            sub_table = Table(sub_data, colWidths=[3*inch, 2*inch, 1.5*inch])
            sub_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), self.BRAND_COLORS["info"]),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                ('BACKGROUND', (0, 1), (-1, -1), colors.white)
            ]))
            elements.append(sub_table)
            elements.append(Spacer(1, 0.15*inch))
        
        return elements
    
    def _create_risk_assessment_section(self, analysis_data: AnalysisData) -> List:
        """Create risk assessment section"""
        elements = []
        
        elements.append(Paragraph("Risk Assessment", self.styles['SectionHeader']))
        elements.append(Spacer(1, 0.2*inch))
        
        risk_assessment = analysis_data.risk_assessment
        
        # Risk summary
        elements.append(Paragraph("Risk Summary", self.styles['SubsectionHeader']))
        elements.append(Paragraph(
            f"Total Risks Identified: {risk_assessment.total_risk_count}",
            self.styles['CustomBody']
        ))
        elements.append(Paragraph(
            f"High Priority: {risk_assessment.high_risk_count} | "
            f"Medium Priority: {risk_assessment.medium_risk_count} | "
            f"Low Priority: {risk_assessment.low_risk_count}",
            self.styles['CustomBody']
        ))
        elements.append(Spacer(1, 0.2*inch))
        
        # Risk details by category
        risks_by_category = {}
        for risk in risk_assessment.risks:
            if risk.category not in risks_by_category:
                risks_by_category[risk.category] = []
            risks_by_category[risk.category].append(risk)
        
        for category, risks in risks_by_category.items():
            elements.append(Paragraph(f"{category} Risks", self.styles['SubsectionHeader']))
            
            for risk in risks:
                # Risk header with severity color
                severity_color = self._get_severity_color(risk.severity)
                elements.append(Paragraph(
                    f"<b>{risk.title}</b> "
                    f"<font color='{severity_color}'>[{risk.severity} Severity, {risk.probability} Probability]</font>",
                    self.styles['CustomBody']
                ))
                
                elements.append(Paragraph(
                    risk.description,
                    self.styles['CustomBody']
                ))
                
                elements.append(Paragraph(
                    "<b>Mitigation Strategies:</b>",
                    self.styles['CustomBody']
                ))
                
                for strategy in risk.mitigation_strategies:
                    elements.append(Paragraph(
                        f"• {strategy}",
                        self.styles['BulletPoint']
                    ))
                
                elements.append(Spacer(1, 0.15*inch))
        
        # Regulatory requirements
        elements.append(Paragraph("Regulatory Requirements", self.styles['SubsectionHeader']))
        for req in risk_assessment.regulatory_requirements:
            elements.append(Paragraph(f"• {req}", self.styles['BulletPoint']))
        
        return elements
    
    def _create_growth_strategy_section(self, analysis_data: AnalysisData) -> List:
        """Create growth strategy section"""
        elements = []
        
        elements.append(Paragraph("Growth Strategy", self.styles['SectionHeader']))
        elements.append(Spacer(1, 0.2*inch))
        
        growth = analysis_data.growth_strategy
        
        # Growth phases
        for phase in growth.growth_phases:
            elements.append(Paragraph(
                f"Phase {phase.phase_number}: {phase.phase_name} ({phase.duration})",
                self.styles['SubsectionHeader']
            ))
            
            elements.append(Paragraph("<b>Objectives:</b>", self.styles['CustomBody']))
            for obj in phase.objectives:
                elements.append(Paragraph(f"• {obj}", self.styles['BulletPoint']))
            
            elements.append(Spacer(1, 0.1*inch))
            elements.append(Paragraph("<b>Key Milestones:</b>", self.styles['CustomBody']))
            for milestone in phase.milestones:
                elements.append(Paragraph(
                    f"• <b>{milestone.title}</b> (Target: {milestone.target_completion}): {milestone.description}",
                    self.styles['BulletPoint']
                ))
            
            elements.append(Spacer(1, 0.2*inch))
        
        # Customer acquisition strategies
        elements.append(Paragraph("Customer Acquisition Strategies", self.styles['SubsectionHeader']))
        
        acq_data = [['Channel', 'Est. Cost', 'Conv. Rate', 'Priority', 'Timeline']]
        for strategy in growth.customer_acquisition_strategies:
            acq_data.append([
                strategy.channel,
                f"${strategy.estimated_cost:,.0f}",
                f"{strategy.expected_conversion_rate:.1f}%",
                strategy.priority,
                strategy.timeline[:20]
            ])
        
        acq_table = Table(acq_data, colWidths=[2*inch, 1.2*inch, 1*inch, 0.8*inch, 2*inch])
        acq_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), self.BRAND_COLORS["accent"]),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
            ('BACKGROUND', (0, 1), (-1, -1), self.BRAND_COLORS["light"]),
            ('VALIGN', (0, 0), (-1, -1), 'TOP')
        ]))
        elements.append(acq_table)
        elements.append(Spacer(1, 0.2*inch))
        
        # 90-day action plan
        elements.append(Paragraph("90-Day Action Plan", self.styles['SubsectionHeader']))
        
        # Group by priority
        high_priority = [a for a in growth.action_plan_90_days if a.priority == "High"]
        
        elements.append(Paragraph("<b>High Priority Actions:</b>", self.styles['CustomBody']))
        for action in high_priority[:10]:  # Limit to top 10
            elements.append(Paragraph(
                f"• <b>{action.title}</b> (Deadline: {action.deadline}, Owner: {action.owner}): {action.description[:80]}...",
                self.styles['BulletPoint']
            ))
        
        return elements
    
    def _create_appendices(self, analysis_data: AnalysisData) -> List:
        """Create appendices section"""
        elements = []
        
        elements.append(Paragraph("Appendices", self.styles['SectionHeader']))
        elements.append(Spacer(1, 0.2*inch))
        
        # Data sources
        elements.append(Paragraph("A. Data Sources & Methodology", self.styles['SubsectionHeader']))
        elements.append(Paragraph(
            "This report was generated using AI-powered analysis combining multiple data sources:",
            self.styles['CustomBody']
        ))
        
        sources = [
            "Market research data from SERP, Apify, and Real Market Service",
            "Industry benchmarks and financial models",
            "AI analysis powered by Groq (Llama 3.3)",
            "Real-time exchange rates for currency conversion"
        ]
        
        for source in sources:
            elements.append(Paragraph(f"• {source}", self.styles['BulletPoint']))
        
        elements.append(Spacer(1, 0.2*inch))
        
        # Disclaimers
        elements.append(Paragraph("B. Disclaimers", self.styles['SubsectionHeader']))
        disclaimers = [
            "This report is for informational purposes only and does not constitute financial, legal, or business advice.",
            "Financial projections are estimates based on industry benchmarks and may not reflect actual results.",
            "Market conditions, competitive landscape, and regulatory requirements may change.",
            "Users should conduct their own due diligence and consult with professional advisors before making business decisions.",
            "AstraMark is not responsible for business outcomes based on this analysis."
        ]
        
        for disclaimer in disclaimers:
            elements.append(Paragraph(f"• {disclaimer}", self.styles['BulletPoint']))
        
        elements.append(Spacer(1, 0.3*inch))
        
        # Report metadata
        elements.append(Paragraph("C. Report Metadata", self.styles['SubsectionHeader']))
        elements.append(Paragraph(
            f"Report ID: {analysis_data.session_id}",
            self.styles['CustomBody']
        ))
        elements.append(Paragraph(
            f"Generated: {analysis_data.generation_timestamp.strftime('%Y-%m-%d %H:%M:%S UTC')}",
            self.styles['CustomBody']
        ))
        elements.append(Paragraph(
            "Platform: AstraMark AI Business Analysis",
            self.styles['CustomBody']
        ))
        
        return elements

    # ==================== DOCX GENERATION METHODS ====================
    
    def _setup_docx_styles(self, doc: Document):
        """Setup custom styles for DOCX document"""
        styles = doc.styles
        
        # Heading 1 style
        if 'Custom Heading 1' not in styles:
            heading1 = styles.add_style('Custom Heading 1', WD_STYLE_TYPE.PARAGRAPH)
            heading1.font.size = Pt(18)
            heading1.font.bold = True
            heading1.font.color.rgb = RGBColor(139, 92, 246)  # Purple
    
    def _add_docx_cover_page(self, doc: Document, analysis_data: AnalysisData):
        """Add cover page to DOCX"""
        # Title
        title = doc.add_heading('Business Feasibility Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Business title
        business_title = analysis_data.business_idea.get('description', 'Business Analysis')[:100]
        subtitle = doc.add_heading(business_title, level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Metadata
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"Report ID: {analysis_data.session_id[:16]}\n").bold = True
        p.add_run(f"Generated: {analysis_data.generation_timestamp.strftime('%B %d, %Y')}\n")
        
        currency_symbol = self._get_currency_symbol(analysis_data.budget.currency)
        p.add_run(f"Budget: {currency_symbol}{analysis_data.budget.total_budget:,.0f}\n")
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        footer = doc.add_paragraph("Generated by AstraMark AI Business Analysis Platform")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_docx_executive_summary(self, doc: Document, analysis_data: AnalysisData):
        """Add executive summary to DOCX"""
        doc.add_heading('Executive Summary', 1)
        
        # Business concept
        doc.add_paragraph(f"Business Concept: {analysis_data.business_idea.get('description', 'N/A')}")
        
        # Key findings
        doc.add_heading('Key Findings', 2)
        
        fp = analysis_data.financial_projections
        year1_revenue = sum(rp.revenue for rp in fp.revenue_projections[:12])
        
        findings = [
            f"Projected Year 1 Revenue: ${year1_revenue:,.0f}",
            f"Break-even Point: Month {fp.break_even_analysis.break_even_month}",
            f"3-Year ROI: {fp.roi_calculations['3_year'].roi_percentage:.1f}%",
            f"Total Risks Identified: {analysis_data.risk_assessment.total_risk_count}",
            f"Growth Phases: {len(analysis_data.growth_strategy.growth_phases)} phases"
        ]
        
        for finding in findings:
            doc.add_paragraph(finding, style='List Bullet')
    
    def _add_docx_business_overview(self, doc: Document, analysis_data: AnalysisData):
        """Add business overview to DOCX"""
        doc.add_heading('Business Overview', 1)
        
        idea = analysis_data.business_idea
        
        doc.add_paragraph(f"Description: {idea.get('description', 'N/A')}")
        doc.add_paragraph(f"Industry: {idea.get('industry', 'N/A')}")
        doc.add_paragraph(f"Target Market: {idea.get('target_market', 'N/A')}")
        doc.add_paragraph(f"Geographic Location: {idea.get('geographic_location', 'N/A')}")
    
    def _add_docx_market_analysis(self, doc: Document, analysis_data: AnalysisData):
        """Add market analysis to DOCX"""
        doc.add_heading('Market Analysis', 1)
        
        market = analysis_data.market_research or {}
        
        if 'market_size' in market:
            doc.add_heading('Market Size', 2)
            doc.add_paragraph(f"Total Addressable Market: {market['market_size'].get('value', 'N/A')}")
        
        if 'competitors' in market:
            doc.add_heading('Competitors', 2)
            doc.add_paragraph(f"Identified {len(market['competitors'])} direct competitors")
    
    def _add_docx_financial_projections(self, doc: Document, analysis_data: AnalysisData):
        """Add financial projections to DOCX"""
        doc.add_heading('Financial Projections', 1)
        
        fp = analysis_data.financial_projections
        
        # Break-even
        doc.add_heading('Break-Even Analysis', 2)
        doc.add_paragraph(f"Break-Even Month: Month {fp.break_even_analysis.break_even_month}")
        doc.add_paragraph(f"Break-Even Revenue: ${fp.break_even_analysis.break_even_revenue:,.0f}")
        
        # ROI
        doc.add_heading('Return on Investment', 2)
        for period_key, roi in fp.roi_calculations.items():
            period_label = period_key.replace('_', ' ').title()
            doc.add_paragraph(f"{period_label}: {roi.roi_percentage:.1f}% ROI", style='List Bullet')
    
    def _add_docx_budget_breakdown(self, doc: Document, analysis_data: AnalysisData):
        """Add budget breakdown to DOCX"""
        doc.add_heading('Budget Breakdown', 1)
        
        budget = analysis_data.budget
        currency_symbol = self._get_currency_symbol(budget.currency)
        
        for cat in budget.categories:
            doc.add_heading(cat.category, 2)
            doc.add_paragraph(f"Amount: {currency_symbol}{cat.amount:,.0f} ({cat.percentage:.1f}%)")
            doc.add_paragraph(f"Rationale: {cat.rationale}")
    
    def _add_docx_risk_assessment(self, doc: Document, analysis_data: AnalysisData):
        """Add risk assessment to DOCX"""
        doc.add_heading('Risk Assessment', 1)
        
        risk_assessment = analysis_data.risk_assessment
        
        doc.add_paragraph(f"Total Risks: {risk_assessment.total_risk_count}")
        doc.add_paragraph(f"High Priority: {risk_assessment.high_risk_count}")
        
        for risk in risk_assessment.risks[:10]:  # Top 10 risks
            doc.add_heading(f"{risk.title} [{risk.severity}]", 2)
            doc.add_paragraph(risk.description)
            doc.add_paragraph("Mitigation Strategies:")
            for strategy in risk.mitigation_strategies:
                doc.add_paragraph(strategy, style='List Bullet')
    
    def _add_docx_growth_strategy(self, doc: Document, analysis_data: AnalysisData):
        """Add growth strategy to DOCX"""
        doc.add_heading('Growth Strategy', 1)
        
        growth = analysis_data.growth_strategy
        
        for phase in growth.growth_phases:
            doc.add_heading(f"Phase {phase.phase_number}: {phase.phase_name}", 2)
            doc.add_paragraph(f"Duration: {phase.duration}")
            doc.add_paragraph("Objectives:")
            for obj in phase.objectives:
                doc.add_paragraph(obj, style='List Bullet')
    
    # ==================== HELPER METHODS ====================
    
    def _generate_report_id(self) -> str:
        """Generate unique report ID"""
        timestamp = datetime.now(timezone.utc).strftime('%Y%m%d%H%M%S')
        random_suffix = secrets.token_hex(4)
        return f"RPT-{timestamp}-{random_suffix}"
    
    def _get_currency_symbol(self, currency: str) -> str:
        """Get currency symbol"""
        symbols = {
            "INR": "₹",
            "USD": "$",
            "EUR": "€"
        }
        return symbols.get(currency.upper(), "$")
    
    def _get_severity_color(self, severity: str) -> str:
        """Get color for risk severity"""
        colors = {
            "High": "#EF4444",
            "Medium": "#F59E0B",
            "Low": "#10B981"
        }
        return colors.get(severity, "#6B7280")
    
    def _add_page_number(self, canvas, doc):
        """Add page numbers to PDF"""
        page_num = canvas.getPageNumber()
        text = f"Page {page_num}"
        canvas.saveState()
        canvas.setFont('Helvetica', 9)
        canvas.drawRightString(7.5*inch, 0.5*inch, text)
        canvas.restoreState()


# Singleton instance
report_generator = ReportGenerator()
